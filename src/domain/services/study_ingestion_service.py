from __future__ import annotations

import base64
import io
import json
import re
import subprocess
import tempfile
from dataclasses import dataclass

import pymupdf
import pytesseract
from PIL import Image

from src.core.settings import Settings
from src.dal.remote.cortex_adapter import CortexAdapter, CortexUnavailableError
from src.dal.remote.fsm_media_adapter import FsmMediaAdapter
from src.domain.models.study import SourceKind, SourceStatus, StudySource


class IngestionError(RuntimeError):
    pass


@dataclass(frozen=True)
class IngestedSource:
    text: str
    chunks: list[dict[str, object]]

    def as_bytes(self) -> bytes:
        return json.dumps({"version": 1, "text": self.text, "chunks": self.chunks}, ensure_ascii=False, separators=(",", ":")).encode()


class StudyIngestionService:
    """Extracts only selected content and persists compact, provenance-bearing text.

    No user token or third-party client is involved. Audio is handed to
    Cortex's existing attachment/Whisper path; deterministic normalisation and
    chunking avoid wasting a higher tier on repeated preparation work.
    """

    def __init__(self, *, fsm: FsmMediaAdapter, cortex: CortexAdapter, settings: Settings) -> None:
        self._fsm = fsm
        self._cortex = cortex
        self._settings = settings

    async def ingest(self, source: StudySource) -> IngestedSource:
        raw = await self._fsm.get(key=source.object_key)
        # When no selection was set the caller wants the whole document.
        # Derive safe full-document bounds from the raw bytes so the rest of
        # the pipeline can proceed without a prior /selection call.
        selection = source.selection
        if selection is None:
            selection = self._full_selection(source.kind, raw)
        if source.kind is SourceKind.pdf:
            text, ranges = self._extract_pdf(raw, selection.page_start or 1, selection.page_end or 1)
        elif source.kind is SourceKind.docx:
            text, ranges = self._extract_docx(raw, selection.page_start or 1, selection.page_end or 1)
        elif source.kind is SourceKind.csv:
            text, ranges = self._extract_csv(raw, selection.line_start or 1, selection.line_end or 1)
        elif source.kind is SourceKind.text:
            text, ranges = self._extract_text(raw, selection.line_start or 1, selection.line_end or 1)
        else:
            text = await self._transcribe_audio(raw, source)
            ranges = [{"audio_start_ms": selection.audio_start_ms or 0, "audio_end_ms": selection.audio_end_ms or 0}]
        normalized = self._normalize(text)
        if not normalized:
            raise IngestionError("The selected source did not yield usable text")
        return IngestedSource(text=normalized, chunks=self._chunk(normalized, ranges))

    @staticmethod
    def _extract_pdf(raw: bytes, start: int, end: int) -> tuple[str, list[dict[str, int]]]:
        try:
            document = pymupdf.open(stream=raw, filetype="pdf")
        except Exception as exc:
            raise IngestionError("The PDF could not be read") from exc
        try:
            if start < 1 or end < start or end > document.page_count:
                raise IngestionError("The selected PDF pages are invalid")
            pages: list[str] = []
            ranges: list[dict[str, int]] = []
            for page_number in range(start, end + 1):
                page = document.load_page(page_number - 1)
                text = page.get_text("text").strip()
                if not text:
                    # OCR only selected image-like pages; never rasterise a
                    # whole PDF unnecessarily.
                    image = Image.open(io.BytesIO(page.get_pixmap(matrix=pymupdf.Matrix(2, 2), alpha=False).tobytes("png")))
                    text = pytesseract.image_to_string(image).strip()
                if text:
                    pages.append(f"[Page {page_number}]\n{text}")
                    ranges.append({"page_start": page_number, "page_end": page_number})
            return "\n\n".join(pages), ranges
        finally:
            document.close()

    @staticmethod
    def _extract_text(raw: bytes, start: int, end: int) -> tuple[str, list[dict[str, int]]]:
        try:
            lines = raw.decode("utf-8-sig").splitlines()
        except UnicodeDecodeError as exc:
            raise IngestionError("Text sources must be UTF-8") from exc
        if start < 1 or end < start or end > len(lines):
            raise IngestionError("The selected text lines are invalid")
        return "\n".join(lines[start - 1:end]), [{"line_start": start, "line_end": end}]

    @staticmethod
    def _extract_docx(raw: bytes, start: int, end: int) -> tuple[str, list[dict[str, int]]]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                raise IngestionError("The Word document has no text content")
            start = max(1, min(start, len(paragraphs)))
            end = min(max(start, end), len(paragraphs))
            selected = paragraphs[start - 1 : end]
            return "\n\n".join(selected), [{"page_start": start, "page_end": end}]
        except Exception as exc:
            raise IngestionError("The Word document could not be read") from exc

    @staticmethod
    def _extract_csv(raw: bytes, start: int, end: int) -> tuple[str, list[dict[str, int]]]:
        try:
            import csv
            content = raw.decode("utf-8-sig")
            reader = list(csv.reader(io.StringIO(content)))
            if not reader:
                raise IngestionError("The CSV file is empty")
            start = max(1, min(start, len(reader)))
            end = min(max(start, end), len(reader))
            selected_rows = [", ".join(row) for row in reader[start - 1 : end]]
            return "\n".join(selected_rows), [{"line_start": start, "line_end": end}]
        except Exception as exc:
            raise IngestionError("The CSV file could not be read") from exc

    async def _transcribe_audio(self, raw: bytes, source: StudySource) -> str:
        # Cortex's documented attachment ingestion invokes its local Whisper
        # capability when enabled. The source remains in FSM; only its text is
        # retained as the compact derivative.
        clipped = self._clip_audio(raw, source)
        attachment = {"filename": "selected-audio.wav", "mime_type": "audio/wav", "data_base64": base64.b64encode(clipped).decode("ascii")}
        prompt = (
            "Transcribe only the selected audio content. Return plain transcript text "
            f"for the interval {source.selection.audio_start_ms}-{source.selection.audio_end_ms} milliseconds."
        )
        try:
            result = await self._cortex.execute_question_generation(prompt=prompt, tier=0, attachments=[attachment])
        except CortexUnavailableError as exc:
            raise IngestionError("Audio transcription is temporarily unavailable") from exc
        if not result.success or not result.response.strip():
            raise IngestionError("Audio transcription failed")
        return result.response

    @staticmethod
    def _clip_audio(raw: bytes, source: StudySource) -> bytes:
        start = source.selection.audio_start_ms or 0
        end = source.selection.audio_end_ms or 0
        if end <= start:
            raise IngestionError("The selected audio interval is invalid")
        try:
            with tempfile.TemporaryDirectory(prefix="certifications-audio-") as directory:
                input_path = f"{directory}/input"
                output_path = f"{directory}/selected.wav"
                with open(input_path, "wb") as handle:
                    handle.write(raw)
                process = subprocess.run(
                    ["ffmpeg", "-v", "error", "-ss", f"{start / 1000:.3f}", "-to", f"{end / 1000:.3f}", "-i", input_path, "-ac", "1", "-ar", "16000", output_path, "-y"],
                    capture_output=True, check=False, timeout=120,
                )
                if process.returncode != 0:
                    raise IngestionError("The selected audio interval could not be read")
                with open(output_path, "rb") as handle:
                    return handle.read()
        except FileNotFoundError as exc:
            raise IngestionError("Audio processing is unavailable") from exc
        except subprocess.TimeoutExpired as exc:
            raise IngestionError("Audio processing timed out") from exc

    @staticmethod
    def _full_selection(kind: SourceKind, raw: bytes) -> "SourceSelection":
        """Derive a whole-document selection when the caller didn't set one."""
        from src.domain.models.study import SourceSelection
        if kind is SourceKind.pdf:
            try:
                doc = pymupdf.open(stream=raw, filetype="pdf")
                pages = doc.page_count
                doc.close()
            except Exception:
                pages = 9999
            return SourceSelection(page_start=1, page_end=pages)
        if kind is SourceKind.docx:
            try:
                import docx
                doc = docx.Document(io.BytesIO(raw))
                paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
            except Exception:
                paragraphs = 9999
            return SourceSelection(page_start=1, page_end=max(1, paragraphs))
        if kind in (SourceKind.csv, SourceKind.text):
            try:
                lines = len(raw.decode("utf-8-sig").splitlines())
            except UnicodeDecodeError:
                lines = 9999
            return SourceSelection(line_start=1, line_end=max(1, lines))
        # audio / video: use a very large end time (10 hours) so all content is covered
        return SourceSelection(audio_start_ms=0, audio_end_ms=10 * 3600 * 1000)

    @staticmethod
    def _normalize(text: str) -> str:
        return re.sub(r"\n{3,}", "\n\n", re.sub(r"[ \t]+", " ", text)).strip()

    @staticmethod
    def _chunk(text: str, source_ranges: list[dict[str, object]], maximum_chars: int = 4000) -> list[dict[str, object]]:
        chunks: list[dict[str, object]] = []
        cursor = 0
        index = 0
        while cursor < len(text):
            end = min(len(text), cursor + maximum_chars)
            if end < len(text):
                boundary = text.rfind("\n", cursor, end)
                if boundary > cursor + maximum_chars // 2:
                    end = boundary
            content = text[cursor:end].strip()
            if content:
                chunks.append({"index": index, "text": content, "source_ranges": source_ranges})
                index += 1
            cursor = max(end, cursor + 1)
        return chunks
